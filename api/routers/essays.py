"""Essay CRUD endpoints — markdown files with YAML frontmatter."""
import io
import json
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

import frontmatter
import markdown as md_lib
from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import Response

from pydantic import BaseModel

from services.sandbox import data_root, atomic_write
from services.file_parser import extract_text
from services.ai_provider import get_provider

router = APIRouter(prefix="/essays", tags=["essays"])


class EssayCreate(BaseModel):
    title: str = "Untitled Essay"
    topic: Optional[str] = None
    thesis: Optional[str] = None
    citation_style: Optional[str] = None
    target_word_count: Optional[int] = None
    instructions: Optional[str] = None
    writing_type: Optional[str] = "essay"
    extra_fields: Optional[dict] = None


class EssayUpdate(BaseModel):
    title: Optional[str] = None
    topic: Optional[str] = None
    thesis: Optional[str] = None
    content: Optional[str] = None  # markdown string
    outline: Optional[list] = None
    profile_id: Optional[str] = None
    citation_style: Optional[str] = None
    target_word_count: Optional[int] = None
    instructions: Optional[str] = None
    writing_type: Optional[str] = None
    extra_fields: Optional[dict] = None


def _essays_dir() -> Path:
    return data_root() / "essays"


def _load_essay(essay_id: str) -> dict:
    path = _essays_dir() / f"{essay_id}.md"
    if not path.exists():
        # Legacy fallback: try .json during migration window
        legacy = _essays_dir() / f"{essay_id}.json"
        if legacy.exists():
            return _load_legacy_json(essay_id, legacy)
        raise HTTPException(status_code=404, detail="Essay not found")

    post = frontmatter.loads(path.read_text(encoding="utf-8"))
    meta = dict(post.metadata)

    # Load outline from sidecar if it exists
    outline_path = _essays_dir() / f"{essay_id}.outline.json"
    outline = []
    if outline_path.exists():
        try:
            outline = json.loads(outline_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            pass

    return {
        "id": essay_id,
        "title": meta.get("title", "Untitled"),
        "topic": meta.get("topic"),
        "thesis": meta.get("thesis"),
        "profile_id": meta.get("profile_id"),
        "citation_style": meta.get("citation_style"),
        "target_word_count": meta.get("target_word_count"),
        "instructions": meta.get("instructions"),
        "writing_type": meta.get("writing_type", "essay"),
        "extra_fields": meta.get("extra_fields"),
        "content": post.content,
        "outline": outline,
        "ai_checks": _load_ai_checks(essay_id),
        "created_at": meta.get("created_at", ""),
        "updated_at": meta.get("updated_at", ""),
    }


def _load_legacy_json(essay_id: str, path: Path) -> dict:
    """Read a legacy .json essay file (migration window fallback)."""
    data = json.loads(path.read_text(encoding="utf-8"))
    return {
        "id": essay_id,
        "title": data.get("title", "Untitled"),
        "topic": data.get("topic"),
        "thesis": data.get("thesis"),
        "profile_id": data.get("profile_id"),
        "citation_style": data.get("citation_style"),
        "target_word_count": data.get("target_word_count"),
        "instructions": data.get("instructions"),
        "writing_type": data.get("writing_type", "essay"),
        "extra_fields": data.get("extra_fields"),
        "content": data.get("content", ""),
        "outline": data.get("outline", []),
        "ai_checks": data.get("ai_checks", []),
        "created_at": data.get("created_at", ""),
        "updated_at": data.get("updated_at", ""),
    }


def _save_essay(essay_id: str, meta: dict, content: str) -> None:
    """Write essay as markdown with YAML frontmatter."""
    post = frontmatter.Post(content)
    for key in ("title", "topic", "thesis", "profile_id", "citation_style", "target_word_count", "instructions", "writing_type", "extra_fields", "created_at", "updated_at"):
        if key in meta and meta[key] is not None:
            post.metadata[key] = meta[key]

    path = _essays_dir() / f"{essay_id}.md"
    atomic_write(path, frontmatter.dumps(post))


def _save_outline(essay_id: str, outline: list) -> None:
    """Write outline sidecar JSON."""
    path = _essays_dir() / f"{essay_id}.outline.json"
    if outline:
        atomic_write(path, json.dumps(outline, indent=2))
    elif path.exists():
        path.unlink()


def _load_ai_checks(essay_id: str) -> list:
    """Load AI detection checks from sidecar JSON."""
    path = _essays_dir() / f"{essay_id}.ai-checks.json"
    if not path.exists():
        return []
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(payload, list):
            return payload
    except json.JSONDecodeError:
        pass
    return []


@router.get("")
async def list_essays():
    essays = []
    for f in sorted(_essays_dir().glob("*.md"), key=lambda p: p.stat().st_mtime, reverse=True):
        try:
            post = frontmatter.loads(f.read_text(encoding="utf-8"))
            meta = post.metadata
            essay_id = f.stem
            essays.append({
                "id": essay_id,
                "title": meta.get("title", "Untitled"),
                "topic": meta.get("topic"),
                "writing_type": meta.get("writing_type", "essay"),
                "word_count": len(post.content.split()) if post.content else 0,
                "updated_at": meta.get("updated_at"),
                "created_at": meta.get("created_at"),
            })
        except Exception:
            continue
    return essays


@router.post("", status_code=201)
async def create_essay(body: EssayCreate):
    essay_id = str(uuid.uuid4())[:8]
    now = datetime.now(timezone.utc).isoformat()
    meta = {
        "title": body.title,
        "topic": body.topic,
        "thesis": body.thesis,
        "profile_id": None,
        "citation_style": body.citation_style,
        "target_word_count": body.target_word_count,
        "instructions": body.instructions,
        "writing_type": body.writing_type,
        "extra_fields": body.extra_fields,
        "created_at": now,
        "updated_at": now,
    }
    content = ""
    _save_essay(essay_id, meta, content)

    return {
        "id": essay_id,
        "title": body.title,
        "topic": body.topic,
        "thesis": body.thesis,
        "content": content,
        "outline": [],
        "ai_checks": [],
        "profile_id": None,
        "citation_style": body.citation_style,
        "target_word_count": body.target_word_count,
        "instructions": body.instructions,
        "writing_type": body.writing_type,
        "extra_fields": body.extra_fields,
        "created_at": now,
        "updated_at": now,
    }


@router.get("/{essay_id}")
async def get_essay(essay_id: str):
    return _load_essay(essay_id)


@router.put("/{essay_id}")
async def update_essay(essay_id: str, body: EssayUpdate):
    essay = _load_essay(essay_id)
    update_data = body.model_dump(exclude_unset=True)

    # Separate outline from the rest
    outline = update_data.pop("outline", None)

    essay.update(update_data)
    essay["updated_at"] = datetime.now(timezone.utc).isoformat()

    content = essay.pop("content", "")
    essay_id_val = essay.pop("id")
    current_outline = essay.pop("outline", [])

    _save_essay(essay_id_val, essay, content)

    if outline is not None:
        _save_outline(essay_id_val, outline)
        current_outline = outline

    return {
        "id": essay_id_val,
        **essay,
        "content": content,
        "outline": current_outline,
    }


@router.delete("/{essay_id}")
async def delete_essay(essay_id: str):
    path = _essays_dir() / f"{essay_id}.md"
    legacy_path = _essays_dir() / f"{essay_id}.json"
    if not path.exists() and not legacy_path.exists():
        raise HTTPException(status_code=404, detail="Essay not found")
    if path.exists():
        path.unlink()
    if legacy_path.exists():
        legacy_path.unlink()
    # Remove sidecar if exists
    outline_path = _essays_dir() / f"{essay_id}.outline.json"
    if outline_path.exists():
        outline_path.unlink()
    checks_path = _essays_dir() / f"{essay_id}.ai-checks.json"
    if checks_path.exists():
        checks_path.unlink()
    return {"deleted": True}


@router.get("/{essay_id}/export")
async def export_essay(essay_id: str, format: str = "md"):
    essay = _load_essay(essay_id)
    title = essay.get("title", "Untitled")
    safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_') or "essay"
    content = essay.get("content", "")

    if format == "md":
        return Response(
            content=content.encode("utf-8"),
            media_type="application/octet-stream",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.md"'},
        )
    elif format == "html":
        html_body = md_lib.markdown(content, extensions=["tables", "fenced_code"])
        html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
  body {{ max-width: 800px; margin: 2rem auto; padding: 0 1rem; font-family: Georgia, serif; line-height: 1.7; color: #222; }}
  h1, h2, h3 {{ font-family: -apple-system, sans-serif; }}
  h1 {{ border-bottom: 2px solid #ddd; padding-bottom: 0.5rem; }}
  blockquote {{ border-left: 3px solid #ccc; margin-left: 0; padding-left: 1rem; color: #555; }}
  code {{ background: #f5f5f5; padding: 0.2em 0.4em; border-radius: 3px; font-size: 0.9em; }}
  pre code {{ display: block; padding: 1rem; overflow-x: auto; }}
  table {{ border-collapse: collapse; width: 100%; }}
  th, td {{ border: 1px solid #ddd; padding: 0.5rem; text-align: left; }}
</style>
</head>
<body>
<h1>{title}</h1>
{html_body}
</body>
</html>"""
        return Response(
            content=html_doc.encode("utf-8"),
            media_type="application/octet-stream",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.html"'},
        )
    elif format == "fountain":
        writing_type = essay.get("writing_type", "essay")
        if writing_type != "screenplay":
            raise HTTPException(status_code=400, detail="Fountain export is only available for screenplay documents")
        # Convert markdown to Fountain format
        lines = content.split("\n")
        fountain_lines = []
        for line in lines:
            stripped = line.strip()
            # ## headings become scene headings (uppercase)
            if stripped.startswith("## "):
                scene = stripped[3:].upper()
                fountain_lines.append("")
                fountain_lines.append(scene)
            elif stripped.startswith("# "):
                fountain_lines.append("")
                fountain_lines.append(f">{stripped[2:].upper()}<")
            elif stripped.startswith("**") and stripped.endswith("**"):
                # Bold text as character name (uppercase)
                name = stripped.strip("*").upper()
                fountain_lines.append("")
                fountain_lines.append(name)
            elif stripped.startswith("_") and stripped.endswith("_"):
                # Italic as parenthetical
                paren = stripped.strip("_")
                fountain_lines.append(f"({paren})")
            elif stripped.startswith("> "):
                # Blockquote as transition
                fountain_lines.append("")
                fountain_lines.append(f">{stripped[2:].upper()}")
            else:
                fountain_lines.append(line)
        fountain_text = "\n".join(fountain_lines)
        return Response(
            content=fountain_text.encode("utf-8"),
            media_type="application/octet-stream",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.fountain"'},
        )
    elif format == "pdf":
        from fpdf import FPDF
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=25)
        # Use system Arial TTF for full unicode support; fall back to core Helvetica
        _font_dir = Path("/System/Library/Fonts/Supplemental")
        if (_font_dir / "Arial.ttf").exists():
            pdf.add_font("body", "", str(_font_dir / "Arial.ttf"))
            pdf.add_font("body", "B", str(_font_dir / "Arial Bold.ttf"))
            pdf.add_font("body", "I", str(_font_dir / "Arial Italic.ttf"))
            _font = "body"
        else:
            _font = "Helvetica"
        pdf.add_page()
        pdf.set_font(_font, "B", 18)
        pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)
        pdf.set_font(_font, size=12)
        for line in content.split("\n"):
            stripped = line.strip()
            pdf.set_x(pdf.l_margin)  # reset cursor after multi_cell
            heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                sizes = {1: 18, 2: 16, 3: 14, 4: 13, 5: 12, 6: 11}
                pdf.ln(4)
                pdf.set_font(_font, "B", sizes.get(level, 12))
                pdf.cell(0, 10, heading_match.group(2), new_x="LMARGIN", new_y="NEXT")
                pdf.set_font(_font, size=12)
                continue
            if stripped.startswith("> "):
                pdf.set_font(_font, "I", 12)
                pdf.multi_cell(0, 6, stripped[2:])
                pdf.set_font(_font, size=12)
                continue
            if not stripped:
                pdf.ln(4)
                continue
            # Strip basic markdown formatting for plain text output
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', stripped)
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            text = re.sub(r'_([^_]+)_', r'\1', text)
            text = re.sub(r'^[-*]\s+', '  - ', text)
            pdf.multi_cell(0, 6, text)
        pdf_buffer = io.BytesIO()
        pdf.output(pdf_buffer)
        pdf_buffer.seek(0)
        return Response(
            content=pdf_buffer.read(),
            media_type="application/octet-stream",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'},
        )
    elif format == "docx":
        from docx import Document as DocxDocument
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = DocxDocument()
        doc.add_heading(title, level=0)
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            # Headings
            heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                doc.add_heading(heading_match.group(2), level=min(level, 4))
                continue
            # Blockquotes
            if stripped.startswith("> "):
                p = doc.add_paragraph(stripped[2:])
                p.runs[0].italic = True
                continue
            # Bullet list items
            if re.match(r'^[-*]\s+', stripped):
                text = re.sub(r'^[-*]\s+', '', stripped)
                doc.add_paragraph(text, style='List Bullet')
                continue
            # Numbered list items
            num_match = re.match(r'^\d+\.\s+(.*)', stripped)
            if num_match:
                doc.add_paragraph(num_match.group(1), style='List Number')
                continue
            # Regular paragraph with inline formatting
            p = doc.add_paragraph()
            # Split on bold (**text**) and italic (*text* or _text_)
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|_[^_]+_)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif (part.startswith("*") and part.endswith("*")) or (part.startswith("_") and part.endswith("_")):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return Response(
            content=docx_buffer.read(),
            media_type="application/octet-stream",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
        )
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {format}. Use md, html, fountain, pdf, or docx.")


ALLOWED_OUTLINE_MIMES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}

_OUTLINE_PARSE_PROMPT = """Parse the following outline into structured sections.

<outline>
{text}
</outline>

Return a JSON array where each element represents one section:
[
  {{
    "title": "Section heading or topic",
    "notes": "Any sub-points, details, or notes under this section"
  }}
]

Rules:
- Each top-level item (numbered, bulleted, or heading) becomes a section
- Sub-items under a section become that section's notes (joined with newlines)
- Preserve the original ordering
- Keep titles concise (the main heading/topic only)
- If the text has no clear structure, create one section per paragraph with the first sentence as the title
- Return ONLY the JSON array, no other text"""


@router.post("/upload-outline")
async def upload_outline(file: UploadFile = File(...)):
    content_type = file.content_type or ""
    if content_type not in ALLOWED_OUTLINE_MIMES:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type: {content_type}. Allowed: .pdf, .docx, .txt",
        )

    raw = await file.read()
    if len(raw) > 10 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large (max 10MB)")

    # Write to temp file for extraction
    import tempfile
    ext = Path(file.filename or "file").suffix or ".txt"
    tmp = Path(tempfile.mktemp(suffix=ext))
    try:
        tmp.write_bytes(raw)
        text = extract_text(tmp, content_type)
    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))
    finally:
        tmp.unlink(missing_ok=True)

    if not text.strip():
        raise HTTPException(status_code=422, detail="No text could be extracted from the file")

    # Use AI to parse outline text into sections
    prompt = _OUTLINE_PARSE_PROMPT.format(text=text[:50_000])
    provider = get_provider()
    try:
        result = await provider.generate(prompt)
    except RuntimeError as e:
        raise HTTPException(status_code=502, detail=str(e))

    try:
        sections = json.loads(result)
    except json.JSONDecodeError:
        raise HTTPException(status_code=502, detail="AI returned invalid JSON. Try again.")

    if not isinstance(sections, list):
        raise HTTPException(status_code=502, detail="AI returned unexpected format. Try again.")

    # Normalize into OutlineSection format with ids
    outline = []
    for s in sections:
        outline.append({
            "id": str(uuid.uuid4())[:8],
            "title": s.get("title", "Untitled"),
            "notes": s.get("notes", ""),
            "evidence": s.get("evidence", ""),
        })

    return {"sections": outline}
